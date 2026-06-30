#!/usr/bin/env python3
"""
硬件需求解析脚本 — 从 xlsx 设备清单或 docx 技术规格书中提取结构化硬件需求。

用法:
    python3 parse_hardware.py input.xlsx                    # xlsx 模式
    python3 parse_hardware.py input.docx                    # docx 模式
    python3 parse_hardware.py input.docx --json-out out.json

输出: JSON 格式的结构化需求清单，包含自动分类（硬件/软件/服务）、
      强制项检测、配置自设计标记、待澄清标记。
"""

import json
import sys
import os
import re
from pathlib import Path

# ============================================================
# 分类关键词
# ============================================================
HARDWARE_KEYWORDS = [
    '服务器', '一体机', '交换机', '路由器', '防火墙', '网闸', 'NVR',
    '录像机', '工作站', '电脑', '触控屏', '大屏', '光模块', '机柜',
    'UPS', 'NFC标签', '传感器', 'PLC', '控制器', '探针',
    '堡垒机', '审计', '态势感知', '上网行为', '准入', '备份一体机',
    'KVM', '工具包', 'AP', 'AC控制器', 'PoE', '路由器',
]

SOFTWARE_KEYWORDS = [
    '企业版', '授权', '许可', 'License', '虚拟化软件',
    '数据库软件', 'Oracle', 'SQL Server', '中间件', '操作系统',
    'Windows Server', '麒麟', '容灾授权', 'EDR', '杀毒软件',
    '软件license', '虚拟化套件', 'FusionSphere',
]

SERVICE_KEYWORDS = [
    '安装调试', '培训', '运维', '驻场', '等保测评', 'SaaS',
    '维保服务', '实施服务',
]

# 配置自设计模式
SELF_DESIGN_PATTERNS = [
    r'自行设计', r'自行拟定', r'自行配置',
    r'承诺性能达标', r'由投标人根据业务需求自行',
    r'具体配置由投标人',
]

# 模糊描述模式
UNCLEAR_PATTERNS = [
    r'主流品牌', r'知名厂商', r'国产优先',
    r'按需', r'预估', r'暂定',
]


def classify(name, spec):
    """分类: hardware | software | service"""
    combined = f"{name} {spec}".lower()
    for kw in SERVICE_KEYWORDS:
        if kw in combined: return 'service'
    for kw in SOFTWARE_KEYWORDS:
        if kw in combined: return 'software'
    for kw in HARDWARE_KEYWORDS:
        if kw in combined: return 'hardware'
    if re.search(r'(CPU|内存|硬盘|接口|端口|吞吐|电源|机架|网口)', combined):
        return 'hardware'
    return 'unknown'


def is_mandatory(text):
    for p in [r'★', r'不低于', r'不小于', r'必须', r'强制']:
        if re.search(p, text): return True
    return False


def is_self_design(text):
    for p in SELF_DESIGN_PATTERNS:
        if re.search(p, text): return True
    return False


def flag_unclear(text):
    return [p for p in UNCLEAR_PATTERNS if re.search(p, text)]


def _add_item(results, name, spec, seq=''):
    cat = classify(name, spec)
    item = {
        'name': name.strip(),
        'spec': spec.strip(),
        'category': cat,
        'mandatory': is_mandatory(f"{name} {spec}"),
        'is_self_design': is_self_design(spec),
        'unclear_flags': flag_unclear(spec),
    }
    key = {'hardware': 'hardware', 'software': 'software', 'service': 'services'}.get(cat, 'hardware')
    results.setdefault(key, []).append(item)
    if item['is_self_design']:
        results.setdefault('self_design_items', []).append(item)
    if item['unclear_flags']:
        results.setdefault('unclear_items', []).append(item)


def parse_xlsx(filepath):
    import openpyxl
    wb = openpyxl.load_workbook(filepath, data_only=True)
    results = {}
    for sname in wb.sheetnames:
        ws = wb[sname]
        hrow = None
        for ri, row in enumerate(ws.iter_rows(values_only=True), 1):
            row_text = ' '.join(str(c) if c else '' for c in row)
            if any(kw in row_text for kw in ['序号', '设备', '规格']):
                hrow = ri; break
        if not hrow: continue
        cname, cspecs = '', []
        for ri, row in enumerate(ws.iter_rows(values_only=True), hrow + 1):
            cells = [str(c).strip() if c else '' for c in row]
            txt = ' '.join(cells)
            if not txt or '合计' in txt: continue
            seq = cells[0] if cells else ''
            nc = cells[1] if len(cells) > 1 else ''
            if nc and not seq.isdigit(): continue
            if seq.isdigit():
                if cname: _add_item(results, cname, '\n'.join(cspecs), seq)
                cname = nc
                cspecs = [cells[2] if len(cells) > 2 else '']
            elif nc:
                cspecs.append(nc)
            elif any(cells):
                extra = ' '.join(c for c in cells[1:] if c)
                if extra: cspecs.append(extra)
        if cname: _add_item(results, cname, '\n'.join(cspecs))
    return _build(results, filepath, 'xlsx')


def parse_docx(filepath):
    """提取 docx 段落和表格供 LLM 分析"""
    import docx
    doc = docx.Document(filepath)
    dk = '|'.join(HARDWARE_KEYWORDS + SOFTWARE_KEYWORDS + SERVICE_KEYWORDS)
    dre = re.compile(rf'(?:{dk})|(?:CPU|吞吐|内存|硬盘|接口|端口|电源|并发)', re.I)
    out = {'paragraphs': [], 'tables': [], 'stats': {
        'para_count': len(doc.paragraphs), 'table_count': len(doc.tables)}}
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if not t: continue
        bold = any(r.bold for r in para.runs if r.text.strip()) if para.runs else False
        if dre.search(t) or bold:
            out['paragraphs'].append({'index': i, 'text': t,
                                       'is_header': bold, 'has_device_info': bool(dre.search(t))})
    for ti, table in enumerate(doc.tables):
        td = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        out['tables'].append({'index': ti, 'rows': len(table.rows),
                               'cols': len(table.columns), 'data': td})
    return out


def _build(results, filepath, stype):
    hw = results.get('hardware', [])
    sw = results.get('software', [])
    sv = results.get('services', [])
    return {
        'source': os.path.basename(filepath), 'type': stype,
        'hardware': hw, 'software': sw, 'services': sv,
        'summary': {
            'total_items': len(hw) + len(sw) + len(sv),
            'hardware_count': len(hw), 'software_count': len(sw),
            'service_count': len(sv),
            'mandatory_count': sum(1 for i in hw if i.get('mandatory')),
            'unclear_count': len(results.get('unclear_items', [])),
            'self_design_count': len(results.get('self_design_items', [])),
        },
        'unclear_items': [{'name': i['name'], 'flags': i['unclear_flags']}
                          for i in results.get('unclear_items', [])],
        'self_design_items': [{'name': i['name'], 'spec': i['spec'][:200]}
                              for i in results.get('self_design_items', [])],
    }


def main():
    if len(sys.argv) < 2:
        print("用法: python3 parse_hardware.py <input.xlsx|input.docx> [--json-out out.json]")
        sys.exit(1)
    fp = sys.argv[1]
    ext = Path(fp).suffix.lower()
    if ext == '.xlsx': result = parse_xlsx(fp)
    elif ext == '.docx': result = parse_docx(fp)
    else: print(f"不支持: {ext}"); sys.exit(1)

    # Print summary
    if ext == '.xlsx':
        s = result['summary']
        print(f"✅ {s['total_items']}项 (硬件{s['hardware_count']}/软件{s['software_count']}/服务{s['service_count']})")
        if result.get('unclear_items'): print(f"⚠ 待澄清: {len(result['unclear_items'])}项")
        if result.get('self_design_items'): print(f"📐 配置自设计: {len(result['self_design_items'])}项")
    else:
        print(f"✅ {result['stats']['para_count']}段落/{result['stats']['table_count']}表格 已提取")

    jo = None
    for i, a in enumerate(sys.argv):
        if a == '--json-out' and i+1 < len(sys.argv): jo = sys.argv[i+1]
    if jo:
        with open(jo, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"📄 → {jo}")
    else:
        print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == '__main__':
    main()
